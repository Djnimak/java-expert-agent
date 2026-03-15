from pathlib import Path
from docx import Document


def format_docx_table(table, table_index: int) -> str:
    """
    Convert a DOCX table into a readable text block for RAG.
    """
    lines: list[str] = [f"[TABLE {table_index}]"]

    for row_index, row in enumerate(table.rows, start=1):
        cell_values: list[str] = []

        for cell in row.cells:
            text = " ".join(cell.text.split()).strip()
            cell_values.append(text if text else "<empty>")

        lines.append(f"Row {row_index}: " + " | ".join(cell_values))

    return "\n".join(lines)


def extract_text_from_docx(docx_path: Path) -> str:
    """
    Extract text from a DOCX file, including paragraphs and tables.
    """
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    document = Document(str(docx_path))

    parts: list[str] = []

    # Paragraphs
    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split()).strip()
        if text:
            parts.append(text)

    # Tables
    for table_index, table in enumerate(document.tables, start=1):
        table_text = format_docx_table(table, table_index)
        if table_text.strip():
            parts.append(table_text)

    return "\n\n".join(parts).strip()


def load_all_docx(docx_folder: Path) -> dict[str, str]:
    """
    Load all DOCX files from a folder.
    Returns dictionary: {filename: extracted_text}
    """
    if not docx_folder.exists():
        raise FileNotFoundError(f"DOCX folder not found: {docx_folder}")

    docx_texts: dict[str, str] = {}

    for docx_file in sorted(docx_folder.glob("*.docx")):
        print(f"Loading DOCX: {docx_file.name}")
        docx_texts[docx_file.name] = extract_text_from_docx(docx_file)

    return docx_texts